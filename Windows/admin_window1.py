from tkinter import ttk
import tkinter as tk
from datetime import datetime
import customtkinter
from DataBase.conn_to_db import connect_to_db
from docx import Document
import os



def is_whole_number(value):
    print(value)
    try:
        value = float(value)
    except ValueError:
        return False
    print(value)
    if value.is_integer():
        return True
    else:
        return False


def get_current_date():
    return datetime.now().strftime("%Y-%m-%d")


def get_current_time():
    return datetime.now().strftime("%H:%M:%S")


def update_time(cur_date_time, cur_date):
    cur_time = get_current_time()
    cur_date_time.configure(text=f"{cur_date} {cur_time}")
    cur_date_time.after(1000, lambda: update_time(cur_date_time, cur_date))


def fetch_table_name():
    conn = connect_to_db()
    cursor = conn.cursor()
    cursor.execute("""
            SELECT table_name
            FROM information_schema.tables
            WHERE table_schema = 'public'
            AND table_type = 'BASE TABLE';
        """)
    tables_info = cursor.fetchall()
    table_names = [table[0] for table in tables_info]
    cursor.close()
    conn.close()
    return table_names


def fetch_table(table_name):
    conn = connect_to_db()
    cursor = conn.cursor()
    query = f"SELECT * FROM {table_name}"
    cursor.execute(query, (table_name))
    table_list = cursor.fetchall()
    headers = [desc[0] for desc in cursor.description]
    cursor.close()
    conn.close()
    return table_list, headers


def create_table_frame(new_window):
    table_frame = customtkinter.CTkFrame(new_window, width=1200, height=400)
    table_frame.grid(row=0, column=0, padx=10, pady=5)
    return table_frame


def apply_style():
    # apply style
    style = ttk.Style()
    style.theme_use('default')
    style.configure("Treeview",
                    background="#2e2e2e",
                    foreground="white",
                    rowheight=20,
                    fieldbackground="#2e2e2e")
    style.configure("Treeview.Heading",
                    background="#1f1f1f",
                    foreground="white",
                    relief="flat")
    style.map("Treeview.Heading",
              background=[('active', '#3d3d3d')])

    style.configure("Vertical.TScrollbar",
                    background="#3d3d3d",
                    troughcolor="#2e2e2e",
                    arrowcolor="white")


def insert_table_in_frame(headers, table_list, table_frame):
    apply_style()

    tree = ttk.Treeview(table_frame, columns=headers, show='headings', style="Treeview")

    scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=tree.yview, style="Vertical.TScrollbar")
    tree.configure(yscrollcommand=scrollbar.set)
    tree.grid(row=1, column=0, padx=10, pady=5, sticky="nsew")
    scrollbar.grid(row=1, column=1, sticky="ns")

    for header in headers:
        tree.heading(header, text=header)
        tree.column(header, width=150, stretch=tk.YES)  # Ensure width is set and fixed

    for row in table_list:
        tree.insert("", "end", values=row)

    return tree


def update_tables_in_frame(table_list, headers, tree):
    apply_style()

    # Clear all existing items in the tree
    for item in tree.get_children():
        tree.delete(item)

    tree["columns"] = headers
    tree["show"] = "headings"

    # Update headers
    for header in headers:
        tree.heading(header, text=header)
        if header:
            tree.column(header, width=160, stretch=tk.YES)

    # Insert new rows
    for row in table_list:
        tree.insert("", "end", values=row)


class WindowAdmin:
    def __init__(self):
        self.new_value = None
        self.buildingMenu = None
        self.materialNewValue2 = None
        self.materialNewValue1 = None
        self.materialNewValue = None
        self.client = None
        self.contract_id = None
        self.errorCCreate = None
        self.windows = {}
        self.client_id = None

    def createNewWindowTables(self, table_name):
        new_window = customtkinter.CTkToplevel()
        new_window.title(f"{table_name}")
        return new_window

    def display_table_data(self, tableName):
        window_key = f"{self, tableName}"
        if window_key in self.windows and self.windows[window_key].winfo_exists():
            self.windows[window_key].focus()
        else:
            new_window = self.createNewWindowTables(tableName)
            self.windows[window_key] = new_window
            table_frame = create_table_frame(new_window)

            tables = fetch_table_name()
            button_frame = customtkinter.CTkFrame(new_window)
            button_frame.grid(row=1, column=0, padx=10, pady=5, sticky="sw")

            optionmenuL = customtkinter.CTkLabel(button_frame, text="Search table: ",
                                                 fg_color="transparent")
            optionmenuL.grid(row=1, column=0, padx=5, pady=5)

            optionmenu = customtkinter.CTkOptionMenu(button_frame, values=[str(id) for id in tables])
            optionmenu.grid(row=1, column=1, padx=5, pady=5)

            def update_table_in_frame():
                table_list, headers = fetch_table(optionmenu.get())

                if hasattr(update_table_in_frame, 'tree'):
                    update_tables_in_frame(table_list, headers, update_table_in_frame.tree)
                else:
                    update_table_in_frame.tree = insert_table_in_frame(headers, table_list, table_frame)

            confirm_option = customtkinter.CTkButton(button_frame, text="Confirm selection",
                                                     command=update_table_in_frame)
            confirm_option.grid(row=2, column=0, padx=5, pady=5)
            new_window.wm_attributes("-topmost", True)

    def findClientByPayinfo(self, payinfo):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = f"SELECT client_name, client_phone, client_payinfo FROM public.client WHERE client_payinfo = %s"
        cursor.execute(query, (payinfo,))
        client = cursor.fetchone()
        cursor.close()
        conn.close()
        return client

    def getClientId(self, payinfo):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = "SELECT client_id FROM public.client WHERE client_payinfo = %s"
        cursor.execute(query, (payinfo,))
        client_id = cursor.fetchone()
        cursor.close()
        conn.close()
        return client_id[0] if client_id else None

    def getClientIdByName(self, name):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = "SELECT client_id FROM public.client WHERE client_name = %s"
        cursor.execute(query, (name,))
        client_id = cursor.fetchone()
        cursor.close()
        conn.close()
        return client_id[0] if client_id else None

    def getBuildingId(self, bName):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = f"SELECT building_id FROM public.building WHERE building_name = %s"
        cursor.execute(query, (bName,))
        building_id = cursor.fetchone()
        cursor.close()
        conn.close()
        return building_id[0] if building_id else None

    def addClientToTable(self, client):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = """
            INSERT INTO public.client (client_name, client_phone, client_payinfo)
            VALUES (%s, %s, %s)
            RETURNING client_id;
        """
        cursor.execute(query, client)
        new_client_id = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()
        return new_client_id

    def addBuildingToTable(self, building):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = """
            INSERT INTO public.building (client_id, building_name, building_area, city, address)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING building_id;
            """
        cursor.execute(query, building)
        new_building_id = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()
        return new_building_id

    def addContract(self, contract):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = """
            INSERT INTO public.contract (client_id, building_id, date_signed, contract_sum,status)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING contract_id;
            """
        cursor.execute(query, contract)
        new_contract_id = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()
        return new_contract_id

    def get_clients_without_accounts(self):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = """
        SELECT c.client_id, c.client_name
        FROM public.client c
        LEFT JOIN public.accounts a ON c.client_id = a.req_id AND a.acc_type = 'user'
        WHERE a.req_id IS NULL;
        """
        cursor.execute(query)
        results = cursor.fetchall()
        cursor.close()
        conn.close()
        return results

    def getAccounts(self):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = f"SELECT login FROM public.accounts"
        cursor.execute(query, ())
        logins = cursor.fetchall()
        cursor.close()
        conn.close()
        return logins if logins else None

    def getSuppliersId(self, name):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = f"SELECT supplier_id FROM public.supplier WHERE supplier_name = %s"
        cursor.execute(query, (name,))
        supplierID = cursor.fetchone()
        cursor.close()
        conn.close()
        return supplierID

    def fetchTableList(self, table_name):
        conn = connect_to_db()
        cursor = conn.cursor()
        query = f"SELECT * FROM {table_name}"
        cursor.execute(query, (table_name))
        table_list = cursor.fetchall()
        cursor.close()
        conn.close()
        return table_list

    def createContractWindow(self, tableName):
        window_key = f"{self, tableName}"
        if window_key in self.windows and self.windows[window_key].winfo_exists():
            self.windows[window_key].focus()
        else:
            new_window = self.createNewWindowTables(tableName)
            self.windows[window_key] = new_window
            client_frame = customtkinter.CTkFrame(new_window)
            client_frame.grid(row=0, column=0, padx=10, pady=5)
            building_frame = customtkinter.CTkFrame(new_window)
            building_frame.grid(row=0, column=3, padx=10, pady=5)

            clLabel = customtkinter.CTkLabel(client_frame, text="Client info:")
            clLabel.grid(row=1, column=0, padx=5, pady=5)

            clName = customtkinter.CTkEntry(client_frame, placeholder_text="client name")
            clName.grid(row=4, column=0, padx=10, pady=5)
            clPhone = customtkinter.CTkEntry(client_frame, placeholder_text="client number")
            clPhone.grid(row=5, column=0, padx=10, pady=5)
            clPayinfo = customtkinter.CTkEntry(client_frame, placeholder_text="client pay info")
            clPayinfo.grid(row=6, column=0, padx=10, pady=5)

            bLabel = customtkinter.CTkLabel(building_frame, text="building info:")
            bLabel.grid(row=5, column=3, padx=5, pady=5)
            bName = customtkinter.CTkEntry(building_frame, placeholder_text="building name")
            bName.grid(row=6, column=3, padx=10, pady=5)
            bArea = customtkinter.CTkEntry(building_frame, placeholder_text="building area")
            bArea.grid(row=7, column=3, padx=10, pady=5)
            bCity = customtkinter.CTkEntry(building_frame, placeholder_text="building city")
            bCity.grid(row=8, column=3, padx=10, pady=5)
            bAddress = customtkinter.CTkEntry(building_frame, placeholder_text="building address")
            bAddress.grid(row=9, column=3, padx=10, pady=5)

            def createCon():
                clNamev = clName.get()
                clPhonev = clPhone.get()
                clPayinfov = clPayinfo.get()
                bNamev = bName.get()
                bAreav = bArea.get()
                bCityv = bCity.get()
                bAddressv = bAddress.get()
                print(clNamev)

                if not clNamev:
                    clName.focus()
                    self.errorCCreate.configure(text="Fill client name.")
                    return
                if not clPhonev:
                    clPhone.focus()
                    self.errorCCreate.configure(text="Fill client phone number.")
                    return
                if not is_whole_number(clPhonev):
                    clPhone.focus()
                    self.errorCCreate.configure(text="Error! Incorrect number format.")
                    return
                if not clPayinfov:
                    clPayinfo.focus()
                    self.errorCCreate.configure(text="Fill client pay info.")
                    return
                if not is_whole_number(clPayinfov):
                    clPayinfo.focus()
                    self.errorCCreate.configure(text="Error! Incorrect pay info format.")
                    return
                if not bNamev:
                    bName.focus()
                    self.errorCCreate.configure(text="Fill building name.")
                    return
                if not bAreav:
                    bArea.focus()
                    self.errorCCreate.configure(text="Fill building area.")
                    return
                if not is_whole_number(bAreav):
                    bArea.focus()
                    self.errorCCreate.configure(text="Error! Incorrect building area format.")
                    return
                if not bCityv:
                    bCity.focus()
                    self.errorCCreate.configure(text="Fill building city.")
                    return
                if not bAddressv:
                    bAddress.focus()
                    self.errorCCreate.configure(text="Fill building address.")
                    return

                client = self.findClientByPayinfo(clPayinfo.get())

                # added client to tables, got client_id
                if client:
                    client_id = self.getClientId(clPayinfo.get())
                    print(client_id)
                else:
                    client = (clName.get(), clPhone.get(), clPayinfo.get())
                    client_id = self.addClientToTable(client)
                    print(client_id)

                building = (client_id, bName.get(), bArea.get(), bCity.get(), bAddress.get())
                building_id = self.addBuildingToTable(building)
                # added building to tables, got building_id
                print(building_id)

                # create contract using variables
                contract = (client_id, building_id, get_current_date(), 0.00, "active")
                self.contract_id = self.addContract(contract)
                if self.contract_id:
                    self.errorCCreate.configure(
                        text=f"Contract successfully created\n Contract id : {self.contract_id}")
                    cPrint.configure(state="normal")
                    cCreate.configure(state="disabled")
                else:
                    self.errorCCreate.configure(text="Contract NOT CREATED")
                    cCreate.configure(state="disabled")

            def printContract():
                # Create a new Document
                doc = Document()

                # Title
                doc.add_heading('Contract', 0)

                # Client Info
                doc.add_heading('Client Info:', level=1)
                doc.add_paragraph(f"Client Name: {clName.get()}", style='BodyText')
                doc.add_paragraph(f"Client Phone Number: {clPhone.get()}", style='BodyText')
                doc.add_paragraph(f"Client Pay Info: {clPayinfo.get()}", style='BodyText')

                # Building Info
                doc.add_heading('Building Info:', level=1)
                doc.add_paragraph(f"Building Name: {bName.get()}", style='BodyText')
                doc.add_paragraph(f"Building Area: {bArea.get()}", style='BodyText')
                doc.add_paragraph(f"Building City: {bCity.get()}", style='BodyText')
                doc.add_paragraph(f"Building Address: {bAddress.get()}", style='BodyText')

                # Contract Info
                doc.add_heading('Contract Info:', level=1)
                current_date = datetime.now().strftime("%Y-%m-%d")
                doc.add_paragraph(f"Contract ID: {self.contract_id}", style='BodyText')
                doc.add_paragraph(f"Date of Creation: {current_date}", style='BodyText')

                # Signature Lines
                doc.add_paragraph('\n\nClients Signature: __________________________\n', style='BodyText')
                doc.add_paragraph('Managers Signature: _________________________', style='BodyText')

                # Save the document
                doc_filename = 'contract.docx'
                doc.save(doc_filename)

                self.errorCCreate.configure(text="Contract document created successfully.")
                cPrint.configure(state="normal")

                # Open the document
                os.startfile(doc_filename)

            acceptButtons = customtkinter.CTkFrame(new_window)
            acceptButtons.grid(row=9, column=0, padx=10, pady=5)

            self.errorCCreate = customtkinter.CTkLabel(acceptButtons, text="")
            self.errorCCreate.grid(row=10, column=0, padx=5, pady=5)

            cCreate = customtkinter.CTkButton(acceptButtons, text="Create Contract",
                                              command=lambda: createCon())
            cCreate.grid(row=20, column=0, padx=10, pady=5)

            cPrint = customtkinter.CTkButton(acceptButtons, text="Print Contract", state="disabled",
                                             command=lambda: printContract())
            cPrint.grid(row=21, column=0, padx=10, pady=5)

            new_window.wm_attributes("-topmost", True)

    def createEditWindow(self, tableName):
        window_key = f"{self, tableName}"
        if window_key in self.windows and self.windows[window_key].winfo_exists():
            self.windows[window_key].focus()
        else:
            new_window = self.createNewWindowTables(tableName)
            self.windows[window_key] = new_window

            selectType = customtkinter.CTkFrame(new_window)
            selectType.grid(row=10, column=0, padx=10, pady=5)

            chooseTypeL = customtkinter.CTkLabel(selectType, text="Choose action type:")
            chooseTypeL.grid(row=1, column=0, padx=10, pady=5)
            chooseType = customtkinter.CTkOptionMenu(selectType, values=["add", "edit", "delete"])
            chooseType.grid(row=1, column=1, padx=5, pady=5)

            chooseTableL = customtkinter.CTkLabel(selectType, text="Choose table:")
            chooseTableL.grid(row=2, column=0, padx=10, pady=5)
            tables = fetch_table_name()
            chooseTable = customtkinter.CTkOptionMenu(selectType, values=tables)
            chooseTable.grid(row=2, column=1, padx=5, pady=5)

            inputInfo = customtkinter.CTkFrame(new_window)
            inputInfo.grid(row=20, column=0, padx=10, pady=5)

            def optionmenu_callback(choice):
                if choice == "user":
                    clients = list(set(self.get_clients_without_accounts()))
                    clients = sorted(clients)
                    clients = list(set(row[1] for row in clients))
                    self.client = customtkinter.CTkOptionMenu(inputInfo, values=[str(id) for id in clients])
                    self.client.grid(row=3, column=0, padx=10, pady=5)
                elif choice == "admin":
                    self.client_id = None

            def optionAccChange(choice):
                if choice == "acc_type":
                    self.new_value = customtkinter.CTkOptionMenu(inputInfo, values=["admin", "user"])
                    self.new_value.grid(row=1, column=0, padx=10, pady=5)
                elif choice == "req_id":
                    clients = list(set(self.fetchTableList("client")))
                    clients = sorted(clients)
                    clients = list(set(row[1] for row in clients))
                    self.new_value = customtkinter.CTkOptionMenu(inputInfo, values=[str(name) for name in clients])
                    self.new_value.grid(row=1, column=0, padx=10, pady=5)
                else:
                    self.new_value = customtkinter.CTkEntry(inputInfo, placeholder_text="new value...")
                    self.new_value.grid(row=1, column=0, padx=10, pady=5)

            def actionOnAccounts(actionType):
                if actionType == "add":
                    login = customtkinter.CTkEntry(inputInfo, placeholder_text="login")
                    login.grid(row=0, column=0, padx=10, pady=5)
                    password = customtkinter.CTkEntry(inputInfo, placeholder_text="password")
                    password.grid(row=1, column=0, padx=10, pady=5)
                    accType = customtkinter.CTkOptionMenu(inputInfo, values=["admin", "user"],
                                                          command=optionmenu_callback)
                    accType.grid(row=2, column=0, padx=10, pady=5)
                    self.client = customtkinter.CTkOptionMenu(inputInfo, values=[])
                    return login, password, accType
                elif actionType == "edit":
                    accounts = list(set((self.getAccounts())))
                    accounts = sorted(accounts)
                    accounts = list(set(row[0] for row in accounts))
                    accLogin = customtkinter.CTkOptionMenu(inputInfo, values=[str(login) for login in accounts])
                    accLogin.grid(row=0, column=0, padx=10, pady=5)
                    accChange = customtkinter.CTkOptionMenu(inputInfo,
                                                            values=["login", "password", "acc_type", "req_id"],
                                                            command=optionAccChange)
                    accChange.grid(row=0, column=1, padx=10, pady=5)

                    return accLogin, accChange, self.new_value

                elif actionType == "delete":
                    accounts = list(set((self.getAccounts())))
                    accounts = sorted(accounts)
                    accounts = list(set(row[0] for row in accounts))
                    accLogin = customtkinter.CTkOptionMenu(inputInfo, values=[str(login) for login in accounts])
                    accLogin.grid(row=0, column=0, padx=10, pady=5)
                    return accLogin

            def actionOnMaterial(actionType):
                if actionType == "add":
                    for widget in inputInfo.winfo_children():
                        widget.destroy()
                    suppliers = list(set(self.fetchTableList("supplier")))
                    suppliers = sorted(suppliers)
                    suppliers = list(set(row[1] for row in suppliers))
                    supplier = customtkinter.CTkOptionMenu(inputInfo, values=[str(name) for name in suppliers])
                    supplier.grid(row=0, column=0, padx=10, pady=5)

                    materialName = customtkinter.CTkEntry(inputInfo, placeholder_text="material name")
                    materialName.grid(row=1, column=0, padx=10, pady=5)
                    materialAmount = customtkinter.CTkEntry(inputInfo, placeholder_text="material amount")
                    materialAmount.grid(row=2, column=0, padx=10, pady=5)
                    materialCost = customtkinter.CTkEntry(inputInfo, placeholder_text="material cost")
                    materialCost.grid(row=3, column=0, padx=10, pady=5)

                    materialUnits = customtkinter.CTkOptionMenu(inputInfo, values=["kilo", "piece", "liter", "m^3"])
                    materialUnits.grid(row=4, column=0, padx=10, pady=5)

                    buildings = sorted(list(set(self.fetchTableList("building"))))
                    buildings = list(set(row[2] for row in buildings))
                    materialBuilding = customtkinter.CTkOptionMenu(inputInfo, values=[str(id) for id in buildings])
                    materialBuilding.grid(row=5, column=0, padx=10, pady=5)
                    return supplier, materialName, materialAmount, materialCost, materialUnits, materialBuilding

                elif actionType == "edit":
                    for widget in inputInfo.winfo_children():
                        widget.destroy()
                    materials = list(set((self.fetchTableList("material"))))
                    materials = sorted(materials)
                    materials = list(set(row[2] for row in materials))
                    materialName = customtkinter.CTkOptionMenu(inputInfo, values=[str(name) for name in materials])
                    materialName.grid(row=0, column=0, padx=10, pady=5)

                    def optionMaterial(choice):
                        if choice == "material_units":
                            self.materialNewValue = None

                            self.materialNewValue = customtkinter.CTkOptionMenu(inputInfo,
                                                                                values=["kilo", "piece", "liter",
                                                                                        "m^3"])
                            self.materialNewValue.grid(row=2, column=0, padx=10, pady=5)
                        elif choice == "building_id":
                            self.materialNewValue = customtkinter.CTkOptionMenu(inputInfo, values=None)

                            building_names = sorted(list(set(self.fetchTableList("building"))))
                            building_names = list(set(row[2] for row in building_names))
                            self.buildingMenu = customtkinter.CTkOptionMenu(inputInfo,
                                                                            values=[str(name) for name in
                                                                                    building_names])
                            self.buildingMenu.grid(row=2, column=0, padx=10, pady=5)

                        else:
                            self.materialNewValue = None
                            self.materialNewValue = customtkinter.CTkEntry(inputInfo, placeholder_text="new value...")
                            self.materialNewValue.grid(row=2, column=0, padx=10, pady=5)

                    materialChange = customtkinter.CTkOptionMenu(inputInfo, values=["material_name", "material_dategot",
                                                                                    "material_amount", "material_cost",
                                                                                    "material_units", "building_id"],
                                                                 command=optionMaterial)
                    materialChange.grid(row=0, column=1, padx=10, pady=5)
                    return materialName, materialChange, self.materialNewValue

                elif actionType == "delete":
                    for widget in inputInfo.winfo_children():
                        widget.destroy()
                    materials = list(set((self.fetchTableList("material"))))
                    materials = sorted(materials)
                    materials = list(set(row[2] for row in materials))
                    materialChoose = customtkinter.CTkOptionMenu(inputInfo, values=[str(name) for name in materials])
                    materialChoose.grid(row=1, column=0, padx=10, pady=5)
                    return materialChoose

            def changeTypebutton():
                if chooseTable.get() == "accounts":
                    if chooseType.get() == "add":
                        login, password, accType = actionOnAccounts("add")
                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionAccount("add", login.get(),
                                                                                                     password.get(),
                                                                                                     accType.get()))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)
                    elif chooseType.get() == "edit":
                        accLogin, accChange, new_value = actionOnAccounts("edit")
                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionAccount("edit",
                                                                                                     accLogin.get(),
                                                                                                     accChange.get(),
                                                                                                     new_value))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)
                    elif chooseType.get() == "delete":
                        accLogin = actionOnAccounts("delete")
                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionAccount("delete",
                                                                                                     accLogin.get()))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)
                if chooseTable.get() == "material":
                    if chooseType.get() == "add":
                        supplier, materialName, materialAmount, materialCost, materialUnits, materialBuilding = actionOnMaterial(
                            "add")
                        supplierChe = supplier.get()
                        supplier_id = self.getSuppliersId(str(supplierChe))
                        building_id = self.getBuildingId(materialBuilding.get())
                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionMaterial("add",
                                                                                                      supplier_id,
                                                                                                      materialName.get(),
                                                                                                      materialAmount.get(),
                                                                                                      materialCost.get(),
                                                                                                      materialUnits.get(),
                                                                                                      building_id))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)
                    elif chooseType.get() == "edit":
                        materialName, materialChange, self.materialNewValue = actionOnMaterial("edit")

                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionMaterial("edit",
                                                                                                      materialName.get(),
                                                                                                      materialChange.get(),
                                                                                                      self.materialNewValue.get(),
                                                                                                      self.buildingMenu))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)
                    elif chooseType.get() == "delete":
                        materialChoose = actionOnMaterial("delete")
                        confirmButton = customtkinter.CTkButton(selectType, text="Confirm action",
                                                                command=lambda: confirmActionMaterial("delete",
                                                                                                      materialChoose.get()))
                        confirmButton.grid(row=5, column=0, padx=10, pady=5)

            def confirmActionAccount(actionType, param1, param2=None, param3=None):
                conn = connect_to_db()
                cursor = conn.cursor()
                if actionType == "add":
                    self.client_id = None if self.client is None else self.getClientIdByName(self.client.get())
                    query = """
                                   INSERT INTO public.accounts (login, password, acc_type, req_id)
                                   VALUES (%s, %s, %s, %s)
                                   RETURNING acc_id;
                                   """
                    cursor.execute(query, (param1, param2, param3, self.client_id))
                    acc_id = cursor.fetchone()[0]
                    confirmation_label.configure(text=f"successfully created account with id: {acc_id}")
                elif actionType == "edit":
                    if param2 == "acc_type":
                        param3 = self.new_value.get()
                    elif param2 == "req_id":
                        param3 = self.getClientIdByName(self.new_value.get())
                    else:
                        param3 = self.new_value.get()

                    query = f"UPDATE public.accounts SET {param2} = %s WHERE login = %s"
                    cursor.execute(query, (param3, param1))
                    confirmation_label.configure(text=f"successfully updated account: {param1}")
                elif actionType == "delete":
                    query = "DELETE FROM public.accounts WHERE login = %s"
                    cursor.execute(query, (param1,))
                    confirmation_label.configure(text=f"successfully deleted account: {param1}")
                conn.commit()
                cursor.close()
                conn.close()

            def confirmActionMaterial(actionType, param1, param2=None, param3=None, param4=None, param5=None,
                                      param6=None):
                conn = connect_to_db()
                cursor = conn.cursor()
                if actionType == "add":
                    query = """
                            INSERT INTO public.material (supplier_id, material_name, material_dategot, material_amount, material_cost, material_units, building_id)
                            VALUES (%s, %s, NOW(), %s, %s, %s, %s)
                            RETURNING material_id;
                            """
                    cursor.execute(query, (param1, param2, param3, param4, param5, param6))
                    material_id = cursor.fetchone()[0]
                    confirmation_label.configure(text=f"Successfully created material with id: {material_id}")
                elif actionType == "edit":
                    if param4:
                        param3 = self.getBuildingId(self.buildingMenu.get())

                    query = f"UPDATE public.material SET {param2} = %s WHERE material_name = %s"
                    cursor.execute(query, (param3, param1))
                    confirmation_label.configure(text=f"Successfully updated material: {param1}")
                elif actionType == "delete":
                    query = "DELETE FROM public.material WHERE material_name = %s"
                    cursor.execute(query, (param1,))
                    confirmation_label.configure(text=f"Successfully deleted material: {param1}")
                conn.commit()
                cursor.close()
                conn.close()

            confirmation_label = customtkinter.CTkLabel(selectType, text="", fg_color="transparent")
            confirmation_label.grid(row=40, column=0, padx=20, pady=10)
            chooseTypeB = customtkinter.CTkButton(selectType, text="Choose type", command=changeTypebutton)
            chooseTypeB.grid(row=3, column=0, padx=10, pady=5)

            new_window.wm_attributes("-topmost", True)

    def displayPrompts(self, tableName):
        window_key = f"{self, tableName}"
        if window_key in self.windows and self.windows[window_key].winfo_exists():
            self.windows[window_key].focus()
        else:
            new_window = self.createNewWindowTables(tableName)
            self.windows[window_key] = new_window
            table_frame = create_table_frame(new_window)

            def update_table_in_frame(query):
                conn = connect_to_db()
                cursor = conn.cursor()
                cursor.execute(query)
                rows = cursor.fetchall()
                headers = [description[0] for description in cursor.description]
                if hasattr(update_table_in_frame, 'tree'):
                    update_tables_in_frame(rows, headers, update_table_in_frame.tree)
                else:
                    update_table_in_frame.tree = insert_table_in_frame(headers, rows, table_frame)
                cursor.close()

            # Individual functions for each SQL query
            def unused_materials():
                query = "SELECT * FROM material WHERE material_id NOT IN (SELECT material_id FROM used_materials)"
                update_table_in_frame(query)

            def residential_complexes():
                query = "SELECT * FROM building WHERE building_name LIKE 'ЖК%'"
                update_table_in_frame(query)

            def workers_hired_this_month():
                query = "SELECT * FROM worker WHERE employment_date BETWEEN '2024-05-01' and '2024-05-31'"
                update_table_in_frame(query)

            def count_workers_hired_this_month():
                query = "SELECT COUNT(*) AS workers_hired_this_month FROM worker WHERE employment_date BETWEEN '2024-05-01' AND '2024-05-31'"
                update_table_in_frame(query)

            def assignments_per_worker():
                query = "SELECT w.worker_id, w.fullname, COUNT(a.ass_id) AS assignment_count FROM worker w LEFT JOIN assignment a ON w.worker_id = a.worker_id GROUP BY w.worker_id, w.fullname ORDER BY assignment_count DESC"
                update_table_in_frame(query)

            def buildings_with_jobs():
                query = "SELECT b.building_name, b.building_id, labor_counts.job_count FROM building b JOIN (SELECT l.building_id, COUNT(*) AS job_count FROM labor l GROUP BY l.building_id) labor_counts ON b.building_id = labor_counts.building_id WHERE labor_counts.job_count >= ANY (SELECT COUNT(*) FROM labor l GROUP BY l.building_id)"
                update_table_in_frame(query)

            def workers_without_assignments_this_month():
                query = "SELECT * FROM worker WHERE worker_id NOT IN (SELECT worker_id FROM assignment WHERE labor_id IN (SELECT labor_id FROM labor WHERE date_start >= '2024-05-01' AND date_start <= '2024-05-31'))"
                update_table_in_frame(query)

            def unused_materials_not_exists():
                query = """
                    SELECT material.*
                    FROM material 
                    WHERE NOT EXISTS (
                        SELECT 1
                        FROM used_materials
                        WHERE material.material_id = used_materials.material_id
                    )
                """
                update_table_in_frame(query)

            def work_hours_comments():
                query = """
                    WITH max_hours AS (
                        SELECT a.worker_id, l.labor_name, a.real_hours, a.required_hours
                        FROM assignment a
                        JOIN labor l ON a.labor_id = l.labor_id
                        ORDER BY a.real_hours DESC
                        LIMIT 1
                    ),
                    all_hours_attended AS (
                        SELECT a.worker_id, l.labor_name, a.real_hours, a.required_hours
                        FROM assignment a
                        JOIN labor l ON a.labor_id = l.labor_id
                        WHERE a.real_hours = a.required_hours
                    ),
                    hours_missed AS (
                        SELECT a.worker_id, l.labor_name, a.real_hours, a.required_hours
                        FROM assignment a
                        JOIN labor l ON a.labor_id = l.labor_id
                        WHERE a.real_hours < a.required_hours
                    )
                    SELECT w.worker_id, max_hours.labor_name, w.fullname, max_hours.real_hours, max_hours.required_hours, 'Найбільша кількість годин' AS comment
                    FROM worker w
                    JOIN max_hours ON w.worker_id = max_hours.worker_id
                    UNION ALL
                    SELECT w.worker_id, all_hours_attended.labor_name, w.fullname, all_hours_attended.real_hours, all_hours_attended.required_hours, 'Відвідував усі години' AS comment
                    FROM worker w
                    JOIN all_hours_attended ON w.worker_id = all_hours_attended.worker_id
                    UNION ALL
                    SELECT w.worker_id, hours_missed.labor_name, w.fullname, hours_missed.real_hours, hours_missed.required_hours, 'Не відвідував необхідні години (прогул)' AS comment
                    FROM worker w
                    JOIN hours_missed ON w.worker_id = hours_missed.worker_id
                    ORDER BY worker_id
                """
                update_table_in_frame(query)

            button_frame = customtkinter.CTkFrame(new_window)
            button_frame.grid(row=1, column=0, padx=10, pady=5, sticky="sw")

            # Button to prompt assigning
            buttons = [
                ("Які матеріали не використовувалися на роботах", unused_materials),
                ("Усі будинки, котрі є Житловими Комплексами", residential_complexes),
                ("Працівники, котрі були найняті цього місяця", workers_hired_this_month),
                ("Кількість працівників найнятих цього місяця", count_workers_hired_this_month),
                ("Кількість робіт призначених кожному співробітнику", assignments_per_worker),
                ("Будівлі у яких є хоча б одна виконана робота", buildings_with_jobs),
                ("Співробітники без призначених робіт цього місяця", workers_without_assignments_this_month),
                ("Які матеріали не використовувалися на роботах (заперечення)", unused_materials_not_exists),
                ("Включення коментарю про відвідування роботи співробітником", work_hours_comments)
            ]
            # Create button for each prompt
            for i, (text, func) in enumerate(buttons):
                button = customtkinter.CTkButton(button_frame, text=text, command=func)
                button.grid(row=i, column=0, padx=10, pady=5)


def open_admin_window1(login, app11):
    def toLoginPage():
        print("pressed button")
        admin_window.destroy()
        admin_window.update()
        app11.deiconify()
    window_admin = WindowAdmin()
    admin_window = customtkinter.CTkToplevel()
    admin_window.title("Admin Dashboard")
    admin_window.geometry("800x600")
    cur_login = login
    app11.withdraw()
    cur_date = get_current_date()
    cur_time = get_current_time()
    hello_frame = customtkinter.CTkFrame(admin_window)
    hello_frame.grid(row=0, column=0, padx=10, pady=5)
    hello_label = customtkinter.CTkLabel(hello_frame, text=f"Welcome, {cur_login}")
    hello_label.grid(row=1, column=0, padx=10, pady=5)
    cur_date_time = customtkinter.CTkLabel(hello_frame, text=f"Current time:  {cur_date} {cur_time}")
    cur_date_time.grid(row=2, column=0, padx=10, pady=5)
    update_time(cur_date_time, cur_date)
    select_frame = customtkinter.CTkFrame(admin_window)
    select_frame.grid(row=3, column=0, padx=10, pady=5)
    select_label_desc = customtkinter.CTkLabel(select_frame, text=f"Show data")
    select_label_desc.grid(row=3, column=0, padx=10, pady=5)
    building_button = customtkinter.CTkButton(select_frame, text="View Tables",
                                              command=lambda: window_admin.display_table_data("View Tables"))
    building_button.grid(row=4, column=0, padx=10, pady=5)

    createContract = customtkinter.CTkButton(select_frame, text="Create contract",
                                             command=lambda: window_admin.createContractWindow("Create contract"))
    createContract.grid(row=5, column=0, padx=10, pady=5)

    editData = customtkinter.CTkButton(select_frame, text="Edit tables data",
                                       command=lambda: window_admin.createEditWindow("Edit Tables"))
    editData.grid(row=6, column=0, padx=10, pady=5)

    sqlPrompts = customtkinter.CTkButton(select_frame, text="SQL Prompts",
                                         command=lambda: window_admin.displayPrompts("SQL Prompts"))
    sqlPrompts.grid(row=7, column=0, padx=10, pady=15)

    toLogin = customtkinter.CTkButton(select_frame, text="To Login Page",
                                      command=toLoginPage)
    toLogin.grid(row=10, column=4, padx=10, pady=5, sticky="ns")
